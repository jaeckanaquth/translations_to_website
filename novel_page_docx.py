from docx import Document
from docx.shared import Inches

from googleDocs_post import upload_files

def page_docs(heading, content):
    
    document = Document()

    document.add_heading(heading, 0)

    document.add_paragraph(content)
    head = heading.replace(" ", "-").replace("(", "").replace(")", "").replace("/","-")
    document.save(head + '.docx')
    upload_files(head + '.docx')
    return head