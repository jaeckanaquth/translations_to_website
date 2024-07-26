from docx import Document
from docx.shared import Inches
from googleDocs_post import upload_files
import glob, os
if glob.glob("config.py"):
    import config
else:
    import github_config as config

def page_docs(heading, content):
    docs_folder = config.novel_name.replace(" ", "-")
    if not os.path.exists(docs_folder):
        os.makedirs(docs_folder)
    document = Document()
    document.add_heading(heading, 0)
    document.add_paragraph(content)
    head = heading.replace(" ", "-").replace("(", "").replace(")", "").replace("/","-")
    file = docs_folder + '/' + head + '.docx'
    document.save(file)
    upload_files(docs_folder, head + '.docx')
    return head